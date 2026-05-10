"""
utils/doc_generator.py — Generate downloadable .docx files from agent output.
"""

from __future__ import annotations

import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ─────────────────────────────────────────────────────────────
ACCENT   = RGBColor(0x1A, 0x73, 0xE8)   # Google-blue
DARK     = RGBColor(0x21, 0x21, 0x21)
GREY     = RGBColor(0x75, 0x75, 0x75)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = ACCENT
    run.font.bold = True


def _add_hrule(doc: Document) -> None:
    """Insert a thin horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A73E8")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_cv_docx(rewritten_cv: str, candidate_name: str = "") -> bytes:
    """Return a .docx byte stream for the tailored CV."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(candidate_name or "Curriculum Vitae")
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Tailored CV — {date.today().strftime('%B %Y')}")
    sub_run.font.size      = Pt(10)
    sub_run.font.color.rgb = GREY
    sub_run.font.italic    = True

    _add_hrule(doc)
    doc.add_paragraph()

    # Body — split on blank lines, detect section headers
    for line in rewritten_cv.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Detect all-caps or title-like section headers
        if stripped.isupper() or (stripped.endswith(":") and len(stripped) < 40):
            _add_heading(doc, stripped.rstrip(":"), level=2)
            _add_hrule(doc)
        elif stripped.startswith("•") or stripped.startswith("-"):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped.lstrip("•- "))
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_cover_letter_docx(cover_letter: str, candidate_name: str = "") -> bytes:
    """Return a .docx byte stream for the cover letter."""
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Header
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = header_p.add_run(candidate_name or "Applicant")
    h_run.font.bold      = True
    h_run.font.size      = Pt(13)
    h_run.font.color.rgb = DARK

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(date.today().strftime("%d %B %Y")).font.color.rgb = GREY

    doc.add_paragraph()
    _add_hrule(doc)
    doc.add_paragraph()

    # Body paragraphs
    for para in cover_letter.split("\n\n"):
        stripped = para.strip()
        if stripped:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
